from __future__ import annotations

import json
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = REPO_ROOT / "docs" / "soft-copyright"
SOFTWARE_NAME = "Miao微信小程序软件"
VERSION = "V1.0.0"
GENERATED_DATE = "2026年6月5日"


def set_run_font(run, ascii_font="微软雅黑", east_asia_font="微软雅黑", size=None, bold=None, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_page(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(f"{SOFTWARE_NAME} {VERSION}  {title}  第 ")
    set_run_font(header_run, size=9, color="666666")
    add_page_number(header)
    suffix = header.add_run(" 页")
    set_run_font(suffix, size=9, color="666666")


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Heading 1", 18, "5A2E18"),
        ("Heading 2", 14, "7B4B2A"),
        ("Heading 3", 12, "5A2E18"),
    ]:
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run(title)
    set_run_font(run, size=24, bold=True, color="5A2E18")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    set_run_font(run2, size=16, bold=True, color="7B4B2A")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(220)
    for text in [
        f"软件名称：{SOFTWARE_NAME}",
        f"版本号：{VERSION}",
        "著作权人：请按申请主体填写",
        "开发完成日期：请按实际完成日期填写",
        f"材料生成日期：{GENERATED_DATE}",
    ]:
        run = p3.add_run(text + "\n")
        set_run_font(run, size=12)


def add_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, val in rows:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val
        for idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10.5, bold=(idx == 0))
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run("● ")
        set_run_font(run, size=10.5, color="7B4B2A")
        run2 = p.add_run(item)
        set_run_font(run2, size=10.5)


def add_numbered(doc: Document, items: list[str]):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        run = p.add_run(f"{index}. ")
        set_run_font(run, size=10.5, bold=True, color="7B4B2A")
        run2 = p.add_run(item)
        set_run_font(run2, size=10.5)


def paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def package_info():
    data = json.loads((REPO_ROOT / "package.json").read_text(encoding="utf-8"))
    return data


def create_manual_docx():
    pkg = package_info()
    doc = Document()
    configure_page(doc, "软件说明书")
    configure_styles(doc)
    add_title(doc, f"{SOFTWARE_NAME}", "软件说明书")
    doc.add_page_break()

    doc.add_heading("一、软件概述", level=1)
    paragraph(
        doc,
        "Miao微信小程序软件是一款基于 Taro 与 React 构建的微信小程序客户端软件，"
        "面向养宠与数字陪伴场景，提供 AI 猫咪形象生成、AI 动作视频生成、日常记录、"
        "时光信件、好友互动、积分与个人资料管理等功能。软件通过微信小程序运行环境"
        "与服务端 API 通信，支持账号登录、本地数据缓存、媒体文件上传、异步同步队列、"
        "内容安全检查和运行诊断能力。"
    )
    add_table(
        doc,
        [
            ("软件名称", SOFTWARE_NAME),
            ("版本号", VERSION),
            ("工程目录", "miao-wechat-mini"),
            ("软件类别", "移动互联网应用软件 / 微信小程序客户端软件"),
            ("开发语言与框架", "TypeScript、React、Taro、Less"),
            ("运行环境", "微信客户端、微信小程序运行时、Taro 4.2.0"),
            ("服务端接口", "通过 HTTPS 调用 /api/v1 系列业务接口"),
            ("代码仓库版本", "以当前 mini 仓库最新代码为准"),
        ],
    )

    doc.add_heading("二、运行环境", level=1)
    add_bullets(
        doc,
        [
            "客户端环境：微信小程序运行环境，支持微信 JSCore、原生 Video、Canvas、文件系统、Storage 与 eventCenter 能力。",
            "构建环境：Node.js、npm、Taro CLI、TypeScript、Vitest，构建目标为 weapp，可按 debug 或正式构建输出 dist 产物。",
            "网络环境：小程序需配置合法服务器域名，通过 HTTPS 访问业务接口、AI 任务接口、上传接口、诊断接口和静态媒体资源。",
            "本地存储：使用 Taro Storage 与微信文件系统保存用户登录态、猫咪信息、日记草稿/媒体缓存、积分记录、同步队列和调试配置。",
        ],
    )

    doc.add_heading("三、系统结构", level=1)
    paragraph(
        doc,
        "软件采用页面层、组件层、服务层、工具适配层和资源层的分层结构。页面层负责用户界面与交互流程；"
        "组件层提供头像、日记卡片、弹窗、分享面板、错误边界和自定义底部导航等复用视图；服务层负责认证、"
        "AI 任务、猫咪生命周期、好友关系、积分、日记与同步队列；工具层封装 HTTP 请求、上传、导航、隐私授权、"
        "二维码绘制、分享卡片生成、平台差异和客户端诊断。"
    )
    add_table(
        doc,
        [
            ("页面模块", "welcome、login、register、home、diary、time-letters、points、profile、generation-progress、upload-material、create-companion 等"),
            ("服务模块", "authService、storage、syncQueue、syncManager、volcanoService、contentSafetyService、friendService、secondaryUnlockService"),
            ("工具模块", "httpAdapter、uploadAdapter、storageAdapter、navigateAdapter、eventAdapter、privacyAuthorization、clientDiagnostics"),
            ("公共组件", "CatAvatar、DiaryCard、ConfirmModal、ShareSheet、ErrorBoundary、FrostedGlassBubble、PawLogo、自定义 tab bar"),
        ],
    )

    doc.add_heading("四、主要功能说明", level=1)
    feature_sections = [
        ("1. 启动与登录", [
            "用户进入小程序后首先进入欢迎页，系统根据本地登录态、当前用户资料和猫咪状态决定跳转到登录、创建猫咪或首页。",
            "登录方式包括账号密码登录、微信登录和手机号登录；注册流程校验用户名、密码和昵称，并在登录成功后持久化 token 与用户资料。",
            "当服务端返回 401 未授权时，HTTP 适配器统一清理登录态并触发 auth:unauthorized 事件，避免旧账号数据继续被使用。"
        ]),
        ("2. AI 猫咪创建与形象生成", [
            "用户可通过上传素材或选择创建流程生成猫咪伙伴，系统保存猫咪名称、品种、毛色、头像、占位图和生成状态。",
            "图生图流程使用 AI 配置中的 provider、imageModel、resolution、seed、promptExtend 等参数，生成猫咪锚定形象图。",
            "生成过程页面会显示阶段性进度、状态文案、错误信息和重试入口，支持在微信运行时缺少 AbortController 时进行兼容降级。"
        ]),
        ("3. AI 图生视频与动作解锁", [
            "视频生成流程由 volcanoService 提交任务并轮询结果，支持文件上传方式提交本地图片，避免 base64 超限。",
            "首页故事视频包含 v1_approach、v2_wait、v3_return、v4_fetch 四段动作；系统优先播放完整四段动作模型，支持后台继续补齐剩余动作。",
            "secondaryUnlockService 用于二级动作解锁，切换底部导航或页面显示状态不应中断服务端后台视频生成任务。"
        ]),
        ("4. 首页互动播放", [
            "首页读取当前猫咪与视频资源，使用微信原生 Video 组件播放故事动作，并通过状态机管理 READY、PLAYING_V1、LOOPING_V2、PLAYING_V3、PLAYING_V4。",
            "首页提供气泡文案、互动积分、播放诊断和错误兜底。视频播放异常时上报 clientDiagnostics，便于线上真机问题排查。",
            "自定义底部导航栏保持浮层布局，确保用户可在首页、日记、时光、积分和我的之间切换。"
        ]),
        ("5. 日常记录与好友动态", [
            "用户可以发布文字、单图、多图或视频日记，系统在发布前执行文本安全和媒体安全检查；安全接口不可用时按业务策略降级处理。",
            "日记媒体优先保存到本地文件系统，保存失败时尝试上传至服务端；仍失败时保留微信临时文件引用并交由同步队列后续处理。",
            "日记支持点赞、评论、删除、分享卡片、朋友圈分享和好友动态展示，多图发布沿用 images 字段保存多张图片。"
        ]),
        ("6. 时光信件", [
            "用户可创建面向猫咪伙伴的时光信件，设置解锁时间，系统保存信件内容、猫咪头像、创建时间与解锁时间。",
            "信件页面根据时间判断可读状态，支持提醒设置和本地数据同步。"
        ]),
        ("7. 积分体系", [
            "积分模块记录签到、在线、互动、兑换等交易流水，支持总积分、当日互动积分、在线分钟数和历史记录展示。",
            "创建或解锁新伙伴时可扣除积分；当生成流程中断时，系统按记录进行退还，避免用户积分丢失。"
        ]),
        ("8. 好友与分享", [
            "好友模块支持扫码添加、二维码邀请、好友日记动态、好友资料和分享入口。",
            "分享服务生成小程序分享信息，日记分享可使用 Canvas 生成分享卡片，提升外部分享识别度。"
        ]),
        ("9. 个人资料与设置", [
            "我的页面展示头像、昵称、陪伴天数、记录瞬间和当前伙伴，可进入个人资料、通知、隐私、密码、反馈等设置页面。",
            "头像上传、昵称修改和密码设置通过服务端接口更新，并在本地存储中同步。"
        ]),
        ("10. 调试与发布防线", [
            "debug 构建可注册后台设置页和诊断页，用于查看 AI 配置、积分辅助、客户端诊断和任务状态。",
            "发布脚本包含类型检查、测试、构建、release scan 和 API contract 检查，降低线上产物与服务端接口不一致的风险。"
        ]),
    ]
    for heading, items in feature_sections:
        doc.add_heading(heading, level=2)
        add_numbered(doc, items)

    doc.add_heading("五、用户操作流程", level=1)
    doc.add_heading("1. 首次使用流程", level=2)
    add_numbered(
        doc,
        [
            "用户打开小程序，系统进入欢迎页并检查本地登录态。",
            "未登录用户进入登录页，可选择微信登录、手机号登录、账号密码登录或注册账号。",
            "登录成功后，若没有猫咪伙伴，系统引导进入创建伙伴或上传素材流程。",
            "用户上传猫咪照片并填写名称，确认后进入 AI 形象生成页面。",
            "形象生成成功后，用户确认解锁动作，进入首页等待或查看猫咪互动视频。",
        ],
    )
    doc.add_heading("2. 日记发布流程", level=2)
    add_numbered(
        doc,
        [
            "用户进入日常记录页，点击发布按钮打开记录弹窗。",
            "输入文字内容，并按需选择图片或视频素材。",
            "系统执行隐私授权、文本安全、媒体安全、媒体持久化和数据保存。",
            "发布成功后刷新我的记录列表，并通过同步队列将数据上传至服务端。",
            "若网络异常，系统保留本地记录和待同步任务，后续自动重试。"
        ],
    )
    doc.add_heading("3. 首页互动流程", level=2)
    add_numbered(
        doc,
        [
            "首页加载当前猫咪和视频资源，若存在四段动作视频则进入故事播放模型。",
            "用户点击猫咪或互动区域后，系统根据播放状态切换视频片段并显示气泡文案。",
            "系统记录互动积分，并在视频播放错误或超时时写入诊断日志。",
            "用户可通过底部导航切换到日记、时光、积分或我的页面，不影响已提交到服务端的后台生成任务。"
        ],
    )

    doc.add_heading("六、数据与接口说明", level=1)
    add_table(
        doc,
        [
            ("认证接口", "/api/v1/auth/register、/api/v1/auth/password-login、/api/v1/auth/wechat-login、/api/v1/auth/phone-login、/api/v1/me"),
            ("AI 接口", "/api/v1/ai/tasks、/api/v1/ai/tasks-file、/api/v1/ai/tasks/:taskId"),
            ("安全接口", "/api/v1/security/text、/api/v1/security/media"),
            ("上传接口", "/api/v1/upload、/api/v1/assets/persist-video"),
            ("业务接口", "/api/v1/cats、/api/v1/diaries、/api/v1/letters、/api/v1/points、/api/v1/friends"),
            ("诊断接口", "/api/v1/client-diagnostics、/api/health"),
        ],
    )
    paragraph(
        doc,
        "客户端统一通过 httpAdapter 和 uploadAdapter 调用服务端接口。请求自动附带 X-Client-Type、X-Client-Version 和 Bearer Token。"
        "服务端返回非 2xx 状态时，适配器将错误包装为统一结构；401 状态会触发统一登出逻辑。"
    )

    doc.add_heading("七、安全与隐私设计", level=1)
    add_bullets(
        doc,
        [
            "登录 token 与当前用户信息分开存储，账号切换或退出登录时清理本地缓存和内存缓存。",
            "日记内容发布前进行文本与媒体安全检测，避免不合规内容进入公开或好友动态链路。",
            "调试后台页面仅在 debug/admin 构建中注册，生产构建默认不暴露危险调试入口。",
            "媒体资源通过微信文件系统和服务端上传接口管理，避免长期直接保存大体积 base64 字符串。",
            "客户端诊断仅记录调试需要的事件、状态、错误与资源类型，不应记录用户密码、token 等敏感字段。"
        ],
    )

    doc.add_heading("八、异常处理与稳定性", level=1)
    add_bullets(
        doc,
        [
            "同步队列持久化待同步任务，支持去重、重试、退避、坏任务隔离和耗尽任务清理。",
            "AI 生成任务支持提交失败重试、轮询状态检查、生成失败提示和积分退还。",
            "日记发布在安全接口或上传接口不可用时具备降级策略，减少用户记录丢失。",
            "首页视频播放具备播放状态机、watchdog 超时诊断和错误状态兜底。",
            "文件系统操作包含保存、读取、删除和临时文件清理，降低本地存储膨胀风险。"
        ],
    )

    doc.add_heading("九、构建与发布", level=1)
    add_table(
        doc,
        [
            ("开发构建", "npm run dev:weapp 或 npm run dev:weapp:debug"),
            ("正式构建", "npm run build:weapp"),
            ("H5 构建", "npm run build:h5"),
            ("类型检查", "npm run lint"),
            ("测试", "npm test"),
            ("发布检查", "npm run release:scan、npm run release:api-contract、npm run release:check"),
            ("主要依赖", ", ".join([f"{k}@{v}" for k, v in pkg.get("dependencies", {}).items()])),
        ],
    )

    doc.add_heading("十、维护说明", level=1)
    paragraph(
        doc,
        "软件后续维护应重点关注微信小程序基础库兼容、Taro 版本升级、AI 服务商模型变更、服务端接口契约、"
        "媒体文件大小限制、内容安全接口可用性、发布产物替换和线上诊断日志。每次正式发布前应执行类型检查、"
        "单元测试、微信小程序构建、API 合约检查和真机关键流程回归。"
    )

    doc.save(OUT_DIR / f"{SOFTWARE_NAME}说明书.docx")


def source_files() -> list[Path]:
    roots = [REPO_ROOT / "src", REPO_ROOT / "config", REPO_ROOT / "scripts"]
    extra = [
        REPO_ROOT / "taro.config.js",
        REPO_ROOT / "weapp.config.js",
        REPO_ROOT / "h5.config.js",
        REPO_ROOT / "babel.config.js",
        REPO_ROOT / "src" / "app.config.ts",
        REPO_ROOT / "src" / "app.tsx",
    ]
    exts = {".ts", ".tsx", ".js", ".jsx"}
    files: list[Path] = []
    for root in roots:
        if not root.exists():
            continue
        for path in sorted(root.rglob("*")):
            if path.suffix not in exts:
                continue
            rel = path.relative_to(REPO_ROOT).as_posix()
            if "__tests__" in rel or ".test." in rel:
                continue
            files.append(path)
    files.extend([p for p in extra if p.exists() and p not in files])
    return sorted(set(files), key=lambda p: p.relative_to(REPO_ROOT).as_posix())


def sanitize_code_line(line: str) -> str:
    if re.search(r"(api[_-]?key|secret|password)\s*[:=]\s*['\"][^'\"]+['\"]", line, re.I):
        return re.sub(r"(['\"])[^'\"]+(['\"])", r"\1***REDACTED***\2", line)
    return line.rstrip("\n").replace("\t", "    ")


def collect_source_lines() -> tuple[list[str], int, int]:
    all_lines: list[str] = []
    for path in source_files():
        rel = path.relative_to(REPO_ROOT).as_posix()
        all_lines.append(f"// ===== file: {rel} =====")
        try:
            for line in path.read_text(encoding="utf-8").splitlines():
                all_lines.append(sanitize_code_line(line))
        except UnicodeDecodeError:
            continue
        all_lines.append(f"// ===== end file: {rel} =====")
        all_lines.append("")
    total = len(all_lines)
    if total > 3000:
        selected = all_lines[:1500] + all_lines[-1500:]
    else:
        selected = all_lines
    return selected, total, len(source_files())


def create_code_docx():
    lines, total_lines, file_count = collect_source_lines()
    doc = Document()
    configure_page(doc, "源代码材料")
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.4)

    chunks = [lines[i:i + 50] for i in range(0, len(lines), 50)]
    for page_index, chunk in enumerate(chunks, 1):
        title = doc.add_paragraph()
        title.paragraph_format.space_after = Pt(2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f"{SOFTWARE_NAME} {VERSION} 源程序 第 {page_index} 页")
        set_run_font(title_run, size=8.5, bold=True, color="444444")
        for line_no, line in enumerate(chunk, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(7.2)
            run = p.add_run(f"{line_no:02d}  {line}")
            set_run_font(run, ascii_font="Consolas", east_asia_font="宋体", size=5.6)
        if page_index < len(chunks):
            doc.add_page_break()
    doc.save(OUT_DIR / f"{SOFTWARE_NAME}源代码材料.docx")

    txt_path = OUT_DIR / f"{SOFTWARE_NAME}源代码材料.txt"
    with txt_path.open("w", encoding="utf-8") as f:
        f.write(f"{SOFTWARE_NAME} {VERSION} 源代码材料\n")
        f.write(f"抽取源码文件数：{file_count}\n")
        f.write(f"源码总行数：{total_lines}\n")
        f.write("提交口径：源码超过3000行时，取前1500行和后1500行；每页50行。\n\n")
        for page_index, chunk in enumerate(chunks, 1):
            f.write(f"===== {SOFTWARE_NAME} {VERSION} 源程序 第 {page_index} 页 =====\n")
            for line_no, line in enumerate(chunk, 1):
                f.write(f"{line_no:02d}  {line}\n")
            f.write("\n")


def create_markdown_and_readme():
    lines, total_lines, file_count = collect_source_lines()
    pages = (len(lines) + 49) // 50
    readme = f"""# Miao 微信小程序软著申请材料

本目录基于 `miao-wechat-mini` 当前代码工程生成，用于准备中国计算机软件著作权登记材料。

## 材料清单

- `{SOFTWARE_NAME}说明书.docx`：软件功能、运行环境、操作流程、数据接口、安全与维护说明。
- `{SOFTWARE_NAME}源代码材料.docx`：源程序鉴别材料，按每页 50 行排版。
- `{SOFTWARE_NAME}源代码材料.txt`：同源代码材料的纯文本备查版。

## 生成口径

- 软件名称：{SOFTWARE_NAME}
- 版本号：{VERSION}
- 源码文件数：{file_count}
- 源码总行数：{total_lines}
- 本次源代码材料页数：{pages}
- 源码超过 3000 行时，按常见软著鉴别材料口径取前 1500 行和后 1500 行；不足时取全部源码。
- 已排除 `node_modules`、`dist`、测试文件、图片资源、历史文档和构建产物。

## 提交前需要人工确认

1. 申请表中的“软件名称”和“版本号”需与材料页眉保持完全一致。
2. “著作权人”“开发完成日期”“首次发表日期”等信息需按实际申请主体填写。
3. 如代理机构要求 A4、页眉、页码或页数有额外格式，请以代理机构/登记系统要求为准。
4. 若涉及商业秘密，可在不超过规则允许范围内做脱敏处理；本次仅对明显密钥类字面量做基础脱敏。
"""
    (OUT_DIR / "README.md").write_text(readme, encoding="utf-8")

    md = f"""# {SOFTWARE_NAME}软件说明书（备查版）

正式提交请优先使用同目录下的 DOCX 文件。

## 软件概述

{SOFTWARE_NAME}是一款基于 Taro 与 React 构建的微信小程序客户端软件，提供 AI 猫咪形象生成、AI 图生视频、日常记录、时光信件、好友互动、积分管理、个人资料设置和调试诊断功能。

## 版本信息

- 软件名称：{SOFTWARE_NAME}
- 版本号：{VERSION}
- 工程目录：miao-wechat-mini
- 技术栈：TypeScript、React、Taro、Less
- 运行环境：微信小程序运行时

## 核心模块

- 启动与登录：欢迎页、账号登录、微信登录、手机号登录、注册、重置密码。
- AI 生成：上传素材、创建伙伴、图生图锚定形象、图生视频动作任务、生成进度与失败恢复。
- 首页互动：四段故事视频播放、动作解锁、积分互动、气泡提示和播放诊断。
- 日常记录：文字、单图、多图、视频日记发布，好友动态，评论、点赞、分享卡片。
- 时光信件：创建延时解锁信件、提醒设置和本地同步。
- 积分系统：签到、在线、互动、兑换、交易流水和异常退还。
- 好友系统：扫码添加、二维码邀请、好友动态和分享。
- 设置与诊断：个人资料、通知、隐私、密码、反馈、后台调试和客户端诊断。
"""
    (OUT_DIR / f"{SOFTWARE_NAME}说明书.md").write_text(md, encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_manual_docx()
    create_code_docx()
    create_markdown_and_readme()
    print(f"Generated materials in {OUT_DIR}")


if __name__ == "__main__":
    main()
